from transformers import AutoConfig, AutoModelForSequenceClassification, AutoTokenizer
import numpy as np
from docx import Document
from fastapi import UploadFile, File
import re
from utils import cut_sent, break_file_to_sentences
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches
import io
import base64
import torch
import os

DOCX_OUTPUT_DIR="."

# 插入表格函数
def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)

async def aigc_check(file: UploadFile,device,model,tokenizer):
    # 记录文本内容,包括句子编号、句子内容、人工判别结果、所在段落
    content=[]
    doc=None

    # 判断是否为空文件
    file.file.seek(0, os.SEEK_END)
    file_size = file.file.tell()
    if file_size == 0:
        return "Empty document. Please upload a valid document."

    file.file.seek(0)
    if file.filename.endswith('.docx'):
        raw_file = await file.read()
        with io.BytesIO(raw_file) as f:
            doc = Document(f)
    elif file.filename.endswith('.txt'):
        raw_file = await file.read()
        text_content = raw_file.decode('utf-8')
        lines = text_content.split('\n')
        doc = Document()  # 创建一个新的Document对象
        for line in lines:
            doc.add_paragraph(line)  # 将每一行文本作为一个段落添加到Document中
    else:
        error_message = "Unsupported file format. Only .docx and .txt files are supported."
        return error_message

    paragraph_count=len(doc.paragraphs)  
    sentence_number=0
    for i in range(paragraph_count):
        # 获取段落内容
        paragraph_text = doc.paragraphs[i].text

        # 分句
        paragraph_text = re.sub(r'\n', '', paragraph_text)

        sentences = cut_sent(paragraph_text)
        for sentence in sentences:
            sentence_number += 1
            content.append({'sentence_number': sentence_number, 'sentence': sentence, 'ai_presentage': 0, 'paragraph': i})

    batch = []
    batch_size = 2  # 每组句子的大小
    total_len = len(content)
    i = 0

    model.to(device)

    while i < total_len:
        if total_len - i > batch_size:
            group = content[i:i + batch_size]  # 每次取出一组句子
            # 将每组句子的文本存储在一个列表中，并添加逗号分隔符
            batch = [item['sentence'] for item in group]
            print(batch)

        # 如果剩余的句子不足一组，将剩余的句子作为最后一组
        else:
            last_group = content[i:]
            batch = [item['sentence'] for item in last_group]
            print(batch)

        inputs = tokenizer.batch_encode_plus(
            batch,
            add_special_tokens=True,
            return_tensors="pt",
            padding="max_length",
            truncation=True,
            max_length=256
        )

        input_ids = inputs["input_ids"].to(device)  # 将输入张量移动到 GPU
        attention_mask = inputs["attention_mask"].to(device)  # 将注意力掩码移动到 GPU

        outputs = model(input_ids=input_ids, attention_mask=attention_mask)
        predictions = outputs.logits.detach().cpu().numpy()  # 将预测结果移动到 CPU

        probabilities = np.exp(predictions) / np.sum(np.exp(predictions), axis=1, keepdims=True)
        ai_probabilities = probabilities[:, 1]  # AI生成文本的概率列表
        for j in range(len(batch)):
            content[i + j]['ai_presentage'] = ai_probabilities[j]
        i += batch_size
        torch.cuda.empty_cache()  # 在每个批处理之后清理 GPU 内存

    # 创建一个新的Word文档
    doc_output = Document()
    output_file_path = DOCX_OUTPUT_DIR + "/" + "AI学习通_AIGC检测报告" + ".docx"
    total_words=0
    slight_ai_words=0
    mediate_ai_words=0
    strong_ai_words=0
    ai_content=[]
    paragraph = doc_output.add_paragraph()

    # 设置默认段落样式的字体样式
    doc_output.styles['Normal'].font.name = u'黑体'
    doc_output.styles['Normal'].font.size = Pt(12)# 设置字体大小为12磅
    doc_output.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')

    # 从提供的句子列表中逐个处理
    for item in content:
        sentence = item['sentence']
        sentence = sentence.rstrip() 

        # 加入总字数
        total_words+= len(sentence)
        ai_percentage = item['ai_presentage']
        paragraph_index = item['paragraph']

        # 获取当前段落或创建新段落
        if paragraph_index + 1 == len(doc_output.paragraphs):
            paragraph = doc_output.paragraphs[paragraph_index]
        else:
            paragraph = doc_output.add_paragraph()


        # 设置段落文本
        if ai_percentage >= 0.9:
            strong_ai_words+=len(sentence)
            ai_content.append(item)
            # 创建一个新的 Run 对象，并将其属性设置为红色
            run = paragraph.add_run(sentence)
            font = run.font
            font.color.rgb = RGBColor(255, 0, 0)  # 设置字体颜色为红色
        elif ai_percentage >= 0.8:
            mediate_ai_words+=len(sentence)
            ai_content.append(item)
            # 创建一个新的 Run 对象，并将其属性设置为橙色
            run = paragraph.add_run(sentence)
            font = run.font
            font.color.rgb = RGBColor(255, 104, 0)
        elif ai_percentage >= 0.7:
            slight_ai_words+=len(sentence)
            ai_content.append(item)
            # 创建一个新的 Run 对象，并将其属性设置为深黄色
            run = paragraph.add_run(sentence)
            font = run.font
            font.color.rgb = RGBColor(254, 204, 0)
        else:
            run = paragraph.add_run(sentence)
            font = run.font
            font.color.rgb = RGBColor(0, 0, 0)

    
    # 展示所有疑似ai句子的表
    table2 = doc_output.add_table(rows=1, cols=4)
    table2.style = 'Table Grid'
    table2.autofit = False

    # 设置表格对齐方式为居中
    table2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 设置列宽
    column_widths = [55, 140, 55, 55]
    for i, width in enumerate(column_widths):
        table2.columns[i].width = Pt(width)
        
    header_cells = table2.rows[0].cells
    header_cells[0].text = '序号'
    header_cells[1].text = '疑似AI生成片段'
    header_cells[2].text = '疑似AI生成字数'
    header_cells[3].text = '疑似AIGC全文占比'

    previous_category = None
    previous_paragraph = None
    previous_sentence_number = None
    index = 0

    print(ai_content)

    for item1 in ai_content:
        current_category = None
        if item1['ai_presentage'] >= 0.9:
            current_category = 1
        elif 0.8 <= item1['ai_presentage'] < 0.9:
            current_category = 2
        elif 0.7 <= item1['ai_presentage'] < 0.8:
            current_category = 3
        
        # 当前段落与上一个段落相同，且当前句子与上一个句子相邻，则将当前句子与上一个句子合并在一个单元格输出
        if previous_category == current_category and previous_paragraph == item1['paragraph'] and int(previous_sentence_number)+1==item1['sentence_number']:
            previous_cells = table2.row_cells(index)
            if current_category == 1:
                run = previous_cells[1].paragraphs[0].add_run(item1['sentence'])
                font = run.font
                font.color.rgb = RGBColor(255, 0, 0)
            elif current_category == 2:
                run = previous_cells[1].paragraphs[0].add_run(item1['sentence'])
                font = run.font
                font.color.rgb = RGBColor(255, 104, 0)
            elif current_category == 3:
                run = previous_cells[1].paragraphs[0].add_run(item1['sentence'])
                font = run.font
                font.color.rgb = RGBColor(255, 204, 0)    
            previous_cells[2].text = str(int(previous_cells[2].text) + len(item1['sentence']))
            previous_cells[3].text = str(round(float(previous_cells[3].text[:-1]) + len(item1['sentence']) / total_words*100, 2))+'%'
        else:
            row_cells = table2.add_row().cells
            index += 1
            row_cells[0].text = str(index)
            if current_category == 1:
                run = row_cells[1].paragraphs[0].add_run(item1['sentence'])
                font = run.font
                font.color.rgb = RGBColor(255, 0, 0)
            elif current_category == 2:
                run = row_cells[1].paragraphs[0].add_run(item1['sentence'])
                font = run.font
                font.color.rgb = RGBColor(255, 104, 0)
            elif current_category == 3:
                run = row_cells[1].paragraphs[0].add_run(item1['sentence'])
                font = run.font
                font.color.rgb = RGBColor(255, 204, 0)    
            row_cells[2].text = str(len(item1['sentence']))
            row_cells[3].text = str(round( len(item1['sentence']) / total_words*100, 2))+'%'

        previous_category = current_category
        previous_paragraph = item1['paragraph']
        previous_sentence_number=item1['sentence_number']

    # 居中对齐第1列、第3列和第4列的内容
    for row in table2.rows:
        row.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row.cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row.cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 上下居中对齐
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[2].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[3].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # 居中对齐第1行第2列单元格的内容
    header_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    header_cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # 展示整体疑似ai信息的表
    table1_data = [
        {'序号': 1, '疑似AIGC程度': '轻度疑似', '疑似AIGC概率': '70%-80%', '疑似AI生成字数': str(slight_ai_words),
        '疑似AIGC章节占比': str(round(slight_ai_words/total_words*100,2))+'%'},
        {'序号': 2, '疑似AIGC程度': '中度疑似', '疑似AIGC概率': '80%-90%', '疑似AI生成字数': str(mediate_ai_words),
        '疑似AIGC章节占比': str(round(mediate_ai_words/total_words*100,2))+'%'},
        {'序号': 3, '疑似AIGC程度': '重度疑似', '疑似AIGC概率': '90%以上', '疑似AI生成字数': str(strong_ai_words),
        '疑似AIGC章节占比': str(round(strong_ai_words/total_words*100,2))+'%'}
    ]

    # 疑似AIGC程度分布表
    table1 = doc_output.add_table(rows=1, cols=5)
    table1.style = 'Table Grid'
    # 设置表格对齐方式为居中
    table1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # 设置表头
    header_cells = table1.rows[0].cells
    header_cells[0].text = '序号'
    header_cells[1].text = '疑似AIGC程度'
    header_cells[2].text = '疑似AIGC概率'
    header_cells[3].text = '疑似AI生成字数'
    header_cells[4].text = '疑似AIGC章节占比'

    # 填充数据
    for row_data in table1_data:
        row_cells = table1.add_row().cells
        row_cells[0].text = str(row_data['序号'])
        
        # 设置疑似AIGC程度单元格的字体颜色
        cell_程度 = row_cells[1]
        cell_程度.text = row_data['疑似AIGC程度']
        font_程度 = cell_程度.paragraphs[0].runs[0].font
        if row_data['疑似AIGC程度'] == '轻度疑似':
            font_程度.color.rgb = RGBColor(254, 204, 0)  # 设置字体颜色为深黄色
        elif row_data['疑似AIGC程度'] == '中度疑似':
            font_程度.color.rgb = RGBColor(255, 165, 0)  # 设置字体颜色为橙色
        elif row_data['疑似AIGC程度'] == '重度疑似':
            font_程度.color.rgb = RGBColor(255, 0, 0)  # 设置字体颜色为红色
        
        row_cells[2].text = row_data['疑似AIGC概率']
        row_cells[3].text = str(row_data['疑似AI生成字数'])
        row_cells[4].text = str(row_data['疑似AIGC章节占比'])

    # 调整列宽
    for column in table1.columns:
        column.width = Pt(70)

    # 居中对齐第1列、第3列和第4列的内容
    for row in table1.rows:
        row.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row.cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row.cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row.cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row.cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 上下居中对齐
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[2].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[3].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[4].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


    # AIGC检测结果表
    table3 = doc_output.add_table(rows=2, cols=5)
    table3.style = 'Table Grid'
    table3.autofit = False

    # 设置表格对齐方式为居中
    table3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 设置表头
    header_cells = table3.rows[0].cells
    header_cells[0].text = '检测字符数'
    header_cells[1].text = '人工书写字数'
    header_cells[2].text = '人工占比'
    header_cells[3].text = '疑似AI生成字数'
    header_cells[4].text = '疑似AIGC全文占比'

    # 设置数据行
    data_cells = table3.rows[1].cells
    data_cells[0].text = str(total_words)
    data_cells[1].text = str(total_words-strong_ai_words-slight_ai_words-mediate_ai_words)
    data_cells[2].text = str(round((total_words-strong_ai_words-slight_ai_words-mediate_ai_words)/total_words*100,2))+'%'
    data_cells[3].text = str(strong_ai_words+slight_ai_words+mediate_ai_words)
    data_cells[4].text = str(round((strong_ai_words+slight_ai_words+mediate_ai_words)/total_words*100,2))+'%'

    # 上下左右居中
    for row in table3.rows:
        row.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row.cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row.cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row.cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row.cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 上下居中对齐
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[2].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[3].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[4].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


    doc_output.paragraphs[0].insert_paragraph_before()

    # 将“疑似AIGC片段汇总”插入开头
    doc_output.paragraphs[0].insert_paragraph_before("疑似AIGC片段汇总")
    # 设置字体样式为黑体
    run = doc_output.paragraphs[0].runs[0]
    run.font.bold = True
    # 设置字体大小
    run.font.size = Pt(14)

    # 把表格2移动到“疑似AIGC片段汇总”后
    move_table_after(table2, doc_output.paragraphs[0])

    doc_output.paragraphs[0].insert_paragraph_before()

    # 将“疑似AIGC程度分布”插入
    doc_output.paragraphs[0].insert_paragraph_before("疑似AIGC程度分布")
    # 设置字体样式为黑体
    run = doc_output.paragraphs[0].runs[0]
    run.font.bold = True
    # 设置字体大小
    run.font.size = Pt(14)

    # 再把表格1移动到标题后
    move_table_after(table1, doc_output.paragraphs[0])

    doc_output.paragraphs[0].insert_paragraph_before()

    # 将“AIGC检测结果”插入
    doc_output.paragraphs[0].insert_paragraph_before("AIGC检测结果")
    # 设置字体样式为黑体
    run = doc_output.paragraphs[0].runs[0]
    run.font.bold = True
    # 设置字体大小
    run.font.size = Pt(14)

    # 再把表格1移动到标题后
    move_table_after(table3, doc_output.paragraphs[0])

    doc_output.paragraphs[0].insert_paragraph_before()
    run=doc_output.paragraphs[0].add_run()
    run.add_picture('static/dragonos.jpg', width=Inches(5.00), height=Inches(2.655))
    doc_output.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 将标题插入开头
    doc_output.paragraphs[1].insert_paragraph_before("AIGC检测报告")
    # 设置段落居中对齐
    doc_output.paragraphs[1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # 设置字体样式为黑体
    run = doc_output.paragraphs[1].runs[0]
    run.font.bold = True
    # 设置字体大小为三号字体（16磅）
    run.font.size = Pt(20)

    # # 在页眉中插入图片
    # header_section = doc_output.sections[0].header
    # header_paragraph = header_section.paragraphs[0]
    # run = header_paragraph.add_run()
    # run.add_picture('first_page.png',width=Inches(4.96), height=Inches(0.8))
    # # 设置页眉居中对齐
    # header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 在末尾增加一段
    last_paragraph = doc_output.add_paragraph()

    # 在该段落中插入图片
    run = last_paragraph.add_run()
    run.add_picture('static/dragonos.jpg', width=Inches(5.00), height=Inches(2.655))

    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc_output.save(output_file_path)

    with open(output_file_path, "rb") as f:
        return base64.b64encode(f.read()).decode()

