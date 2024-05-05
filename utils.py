from docx import Document
from fastapi import UploadFile
import re



def cut_sent(para):
    para = re.sub('([。！？\?])([^”’])', r"\1\n\2", para)  # 单字符断句符
    para = re.sub('(\.{6})([^”’])', r"\1\n\2", para)  # 英文省略号
    para = re.sub('(\…{2})([^”’])', r"\1\n\2", para)  # 中文省略号
    para = re.sub('([。！？\?][”’])([^，。！？\?])', r'\1\n\2', para)
    para = para.rstrip()  # 段尾如果有多余的\n就去掉它
    sentences = para.split("\n")
    # sentences = [sentence for sentence in sentences if sentence.strip()]
    return sentences


def break_file_to_sentences(file: UploadFile):
	content=[]
	if file.filename.endswith('.docx'):
		# 读取 UploadFile 文件对象的内容为二进制数据
		file_content = file.read()

		# 使用 docx 的 Document 函数处理文件内容
		doc = Document(file_content)
		paragraph_count=len(doc.paragraphs)

		for i in range(paragraph_count):
			# 获取段落内容
			paragraph_text = doc.paragraphs[i].text

			# 分句
			paragraph_text = re.sub(r'\n', '', paragraph_text)
			sentences = cut_sent(paragraph_text)

			for sentence in sentences:
				content.append({
					'sentence': sentence,
					'ai_presentage': 0,
					'paragraph': i})
	
	elif file.filename.endswith('.txt'):
		# 读取 UploadFile 文件对象的内容为文本数据
		file_content = file.read()

        # 分句
		sentences = cut_sent(file_content)
		
		for i, sentence in enumerate(sentences):
			content.append({
                'sentence': sentence,
                'ai_presentage': 0,
                'paragraph': i
				})

	return content